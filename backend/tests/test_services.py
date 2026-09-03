import asyncio
from datetime import UTC, date, datetime
from types import SimpleNamespace
from unittest.mock import AsyncMock

import app.services.document_service as document_service_module
import app.services.generated_document_service as generated_document_service_module
import app.services.ocr_service as ocr_service_module
import pymupdf
import pytest
from docx import Document as DocxDocument
from app.bot.handlers.cases import _case_detail_text
from app.bot.handlers.cases import create_case
from app.bot.handlers.documents import _save_document
import app.bot.handlers.documents as document_handlers
from app.bot.handlers.legal import _result_text, answer_date_question
from app.bot.states import LegalQuestionnaire
from app.config import get_settings
from app.db.models import (
    Case,
    CaseStatus,
    Document,
    DocumentRecognition,
    GeneratedDocumentType,
    LegalAnalysisStatus,
    LegalAssessment,
    LegalAssessmentStatus,
    LegalGroundStatus,
    RecognitionStatus,
    User,
    UserConsent,
)
from app.db.session import async_session_factory
from app.services.case_service import CaseService
from app.services.consent_service import (
    PERSONAL_DATA_CONSENT_VERSION,
    ConsentService,
)
from app.services.document_service import DocumentService
from app.services.extraction_service import FineNoticeExtractor, FineNoticeFields
from app.services.legal_assessment_service import LegalAssessmentService
from app.services.legal_analysis_service import LegalAnalysisService
from app.services.legal_rules import (
    EvidenceStatus,
    evaluate_rules,
    get_next_question,
    select_rules_version_for_date,
)
from app.services.ocr_service import (
    DisabledOcrProvider,
    OcrResult,
    OcrSpaceProvider,
)
from app.services.recognition_service import RecognitionService
from app.services.user_service import UserService
from sqlalchemy import delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession


class FakeDeepSeekClient:
    def __init__(self) -> None:
        self.document_prompts: list[str] = []

    async def complete_json(self, *, system_prompt, user_prompt, schema):
        from app.ai.schemas import GeneratedLegalDocument, LegalAnalysisResult

        if schema is LegalAnalysisResult:
            return LegalAnalysisResult(
                summary="Есть основания для проверки доказательств.",
                overall_assessment="Недостаточно документов — необходимо истребование доказательств.",
                missing_evidence=["сведения о поверке комплекса"],
                additional_questions=["Есть исходные фотофиксации?"],
                grounds=[
                    {
                        "id": "ground-valid",
                        "title": "Проверка поверки комплекса",
                        "description": "Нет сведений о поверке на дату фиксации.",
                        "supporting_fact_ids": ["article", "answer_camera"],
                        "legal_rule_ids": ["A12"],
                        "source_ids": ["koap-rf"],
                        "missing_evidence": ["сведения о поверке комплекса"],
                        "assumptions": ["данные поверки не представлены"],
                        "recommended": True,
                    },
                    {
                        "id": "ground-invalid-rule",
                        "title": "Нельзя сохранять",
                        "description": "Ссылка на несуществующее правило.",
                        "supporting_fact_ids": ["article"],
                        "legal_rule_ids": ["DOES_NOT_EXIST"],
                        "source_ids": ["koap-rf"],
                    },
                    {
                        "id": "ground-invalid-source",
                        "title": "Неизвестная судебная практика",
                        "description": "Ссылка на неизвестный источник.",
                        "supporting_fact_ids": ["article"],
                        "legal_rule_ids": ["A12"],
                        "source_ids": ["unknown-case"],
                    },
                    {
                        "id": "ground-invalid-fact",
                        "title": "Выдуманная характеристика камеры",
                        "description": "Ссылка на факт, которого нет во входных данных.",
                        "supporting_fact_ids": ["camera_serial_number"],
                        "legal_rule_ids": ["A12"],
                        "source_ids": ["koap-rf"],
                    },
                ],
            )

        assert schema is GeneratedLegalDocument
        self.document_prompts.append(user_prompt)
        assert "Отклоненное основание" not in user_prompt
        return GeneratedLegalDocument(
            title="Жалоба на постановление",
            sections=[
                "Прошу отменить постановление по делу о штрафе.",
                "Используется только подтвержденное основание: Проверка поверки комплекса.",
            ],
        )


@pytest.mark.asyncio
async def test_user_registration_is_idempotent(db_session: AsyncSession) -> None:
    service = UserService(db_session)

    first = await service.get_or_create(100_000_001, "first", "Test", None)
    second = await service.get_or_create(100_000_001, "updated", "Test", "User")

    count = await db_session.scalar(
        select(func.count(User.id)).where(User.telegram_id == 100_000_001)
    )
    assert first.id == second.id
    assert second.username == "updated"
    assert count == 1


@pytest.mark.asyncio
async def test_delete_user_removes_account_data(
    db_session: AsyncSession, tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_011, "delete_test", "Delete", "User"
    )
    case = await CaseService(db_session).create(user.id)
    relative_path = "storage/cases/delete-user/document.pdf"
    file_path = tmp_path / relative_path
    file_path.parent.mkdir(parents=True)
    file_path.write_bytes(b"personal document")
    await DocumentService(db_session).create(
        case=case,
        telegram_file_id="delete-user-file",
        original_filename="document.pdf",
        mime_type="application/pdf",
        local_path=relative_path,
    )
    await ConsentService(db_session).accept_current(user)
    monkeypatch.setattr(document_service_module, "BACKEND_ROOT", tmp_path)

    deleted = await UserService(db_session).delete_by_telegram_id(user.telegram_id)

    assert deleted is True
    assert await UserService(db_session).get_by_telegram_id(user.telegram_id) is None
    assert (
        await db_session.scalar(
            select(func.count(UserConsent.id)).where(UserConsent.user_id == user.id)
        )
        == 0
    )
    assert (
        await db_session.scalar(select(func.count(User.id)).where(User.id == user.id))
        == 0
    )
    assert await CaseService(db_session).get_for_user(user.id, case.id) is None
    assert not file_path.exists()


@pytest.mark.asyncio
async def test_consent_accept_current_is_idempotent(db_session: AsyncSession) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_010, "consent_test", "Consent", "User"
    )
    service = ConsentService(db_session)

    first = await service.accept_current(user)
    second = await service.accept_current(user)

    count = await db_session.scalar(
        select(func.count(UserConsent.id)).where(UserConsent.user_id == user.id)
    )
    latest = await service.latest_for_user(user.id)

    assert first.id == second.id
    assert count == 1
    assert latest is not None
    assert latest.telegram_id == user.telegram_id
    assert latest.version == PERSONAL_DATA_CONSENT_VERSION
    assert await service.has_current_consent(user.id) is True


@pytest.mark.asyncio
async def test_concurrent_user_registration_is_idempotent() -> None:
    telegram_id = 9_999_999_999_999

    async def register(username: str) -> int:
        async with async_session_factory() as session:
            user = await UserService(session).get_or_create(
                telegram_id, username, "Concurrent", "User"
            )
            await session.commit()
            return user.id

    try:
        first_id, second_id = await asyncio.gather(
            register("first"), register("second")
        )
        async with async_session_factory() as session:
            count = await session.scalar(
                select(func.count(User.id)).where(User.telegram_id == telegram_id)
            )

        assert first_id == second_id
        assert count == 1
    finally:
        async with async_session_factory() as session:
            await session.execute(delete(User).where(User.telegram_id == telegram_id))
            await session.commit()


@pytest.mark.asyncio
async def test_case_is_linked_to_user(db_session: AsyncSession) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_002, None, "Case", "Owner"
    )

    case = await CaseService(db_session).create(user.id)
    cases = await CaseService(db_session).list_for_user(user.id)

    assert case.user_id == user.id
    assert [item.id for item in cases] == [case.id]


@pytest.mark.asyncio
async def test_document_is_linked_to_case(db_session: AsyncSession) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_003, None, "Document", "Owner"
    )
    case = await CaseService(db_session).create(user.id)

    document = await DocumentService(db_session).create(
        case=case,
        telegram_file_id="test-file-id",
        original_filename="synthetic.pdf",
        mime_type="application/pdf",
        local_path="storage/cases/test/synthetic.pdf",
    )
    loaded_case = await CaseService(db_session).get_for_user(user.id, case.id)

    assert document.case_id == case.id
    assert loaded_case is not None
    assert loaded_case.status is CaseStatus.DOCUMENT_UPLOADED
    assert [item.id for item in loaded_case.documents] == [document.id]


@pytest.mark.asyncio
async def test_document_upload_starts_required_questionnaire(
    db_session: AsyncSession, monkeypatch: pytest.MonkeyPatch
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_032, None, "Upload", "Questionnaire"
    )
    await ConsentService(db_session).accept_current(user)
    monkeypatch.setattr(
        document_handlers,
        "get_settings",
        lambda: SimpleNamespace(document_storage="telegram", ocr_provider="none"),
    )

    message = SimpleNamespace(
        from_user=SimpleNamespace(id=user.telegram_id),
        answer=AsyncMock(),
    )
    state = AsyncMock()

    await _save_document(
        message=message,
        state=state,
        session=db_session,
        telegram_file_id="telegram-file-id",
        original_filename="notice.pdf",
        mime_type="application/pdf",
    )

    cases = await CaseService(db_session).list_for_user(user.id)
    assessment = await LegalAssessmentService(db_session).get_for_case(cases[0].id)

    assert assessment is not None
    assert assessment.status.value == "IN_PROGRESS"
    assert state.set_state.await_args.args[0] == LegalQuestionnaire.waiting_for_answer
    assert message.answer.await_count == 2
    assert "Дело №" in message.answer.await_args_list[0].args[0]
    assert "Когда вы получили постановление" in message.answer.await_args_list[1].args[0]


@pytest.mark.asyncio
async def test_additional_document_upload_uses_existing_case_without_recognition(
    db_session: AsyncSession, monkeypatch: pytest.MonkeyPatch
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_033, None, "Evidence", "Upload"
    )
    await ConsentService(db_session).accept_current(user)
    case = await CaseService(db_session).create(user.id)
    case.status = CaseStatus.READY
    monkeypatch.setattr(
        document_handlers,
        "get_settings",
        lambda: SimpleNamespace(document_storage="telegram", ocr_provider="none"),
    )

    message = SimpleNamespace(
        from_user=SimpleNamespace(id=user.telegram_id),
        answer=AsyncMock(),
    )
    state = AsyncMock()

    await _save_document(
        message=message,
        state=state,
        session=db_session,
        telegram_file_id="evidence-file-id",
        original_filename="evidence.jpg",
        mime_type="image/jpeg",
        case_id=case.id,
    )

    cases = await CaseService(db_session).list_for_user(user.id)
    loaded_case = await CaseService(db_session).get_for_user(user.id, case.id)

    assert [item.id for item in cases] == [case.id]
    assert loaded_case is not None
    assert loaded_case.status is CaseStatus.READY
    document = await db_session.scalar(
        select(Document).where(Document.case_id == case.id)
    )
    assert document is not None
    recognitions_count = await db_session.scalar(
        select(func.count(DocumentRecognition.id)).where(
            DocumentRecognition.document_id == document.id
        )
    )
    assert recognitions_count == 0
    assert "Материалы добавлены" in message.answer.await_args.args[0]


def test_legal_rules_derive_directions_and_evidence_from_answers() -> None:
    results = evaluate_rules(
        {
            "appeal_received_at": "28.08.2026",
            "complaint_recipient": "Тверской районный суд города Москвы",
            "correspondence_address": "Москва, Тверская улица, дом 1",
            "driver": "other",
            "driver_docs": "yes",
            "vehicle_photo": "different",
            "plate_photo": "different",
            "place_time_match": "yes",
            "speed": "dispute",
            "speed_docs": "no",
            "camera": "calibration",
            "sign": "hidden",
            "sign_docs": "yes",
            "marking": "none",
            "owner_data_match": "yes",
            "previous_resolution": "no",
            "article_qualification": "no",
            "duplicate": "yes",
            "duplicate_docs": "yes",
            "emergency": "yes",
            "emergency_docs": "no",
        }
    )

    by_code = {result["code"]: result for result in results}
    assert set(by_code) == {"A01", "A04", "A05", "A07", "A09", "A12", "A16", "A17"}
    assert by_code["A01"]["evidence_status"] == EvidenceStatus.AVAILABLE.value
    assert by_code["A07"]["evidence_status"] == EvidenceStatus.NEEDED.value
    assert by_code["A12"]["evidence_status"] == EvidenceStatus.VERIFY.value
    assert by_code["A01"]["evidence_items"][0]["status"] == EvidenceStatus.AVAILABLE.value


def test_extended_stage_three_scenarios_are_evaluated() -> None:
    results = evaluate_rules(
        {
            "appeal_received_at": "28.08.2026",
            "complaint_recipient": "Тверской районный суд города Москвы",
            "correspondence_address": "Москва, Тверская улица, дом 1",
            "driver": "lost",
            "possession_docs": "yes",
            "vehicle_photo": "no_photo",
            "plate_photo": "no_photo",
            "place_time_match": "wrong_place",
            "place_time_docs": "yes",
            "camera": "none",
            "sign": "none",
            "marking": "conflict",
            "marking_docs": "no",
            "owner_data_match": "wrong",
            "owner_data_docs": "yes",
            "previous_resolution": "cancelled",
            "previous_resolution_docs": "yes",
            "article_qualification": "yes",
            "duplicate": "no",
            "emergency": "no",
        },
        notice_article="ч.1 ст.12.16",
    )

    by_code = {result["code"]: result for result in results}
    assert set(by_code) == {"A03", "A06", "A08", "A10", "A14", "A15", "A18"}
    assert by_code["A03"]["evidence_status"] == EvidenceStatus.AVAILABLE.value
    assert by_code["A06"]["evidence_status"] == EvidenceStatus.VERIFY.value
    assert by_code["A10"]["evidence_status"] == EvidenceStatus.NEEDED.value
    assert by_code["A18"]["evidence_status"] == EvidenceStatus.VERIFY.value


def test_legal_result_explains_evidence_without_empty_technical_labels() -> None:
    assessment = LegalAssessment(
        answers={"appeal_received_at": "28.08.2026"},
        results=[
            {
                "code": "A10",
                "title": "Дорожная разметка отсутствовала или была не видна",
                "direction": "Проверить состояние разметки на участке нарушения.",
                "reasons": ["Вы указали, что разметка отсутствовала."],
                "evidence_items": [
                    {
                        "name": "Фото или видео участка",
                        "status": EvidenceStatus.NEEDED.value,
                    }
                ],
            }
        ],
    )

    text = _result_text(assessment)

    assert text.startswith(
        "Результат предварительной проверки оснований для обжалования"
    )
    assert "Что проверить:" in text
    assert "Почему выбрано это направление:" in text
    assert "Для подтверждения понадобятся: Фото или видео участка." in text
    assert "Есть: нет" not in text
    assert "Нужно запросить: нет" not in text
    assert "Следующий шаг: запустите AI-анализ" in text

FIRST_TEN_SCENARIO_CASES = (
    ("A01", {"driver": "other", "driver_docs": "yes"}, {"driver": "owner"}, {"A02", "A03"}),
    ("A02", {"driver": "sold", "sale_docs": "yes"}, {"driver": "owner"}, {"A01", "A03"}),
    ("A03", {"driver": "lost", "possession_docs": "yes"}, {"driver": "owner"}, {"A01", "A02"}),
    ("A04", {"plate_photo": "different"}, {"plate_photo": "yes"}, set()),
    ("A05", {"vehicle_photo": "different"}, {"vehicle_photo": "yes"}, {"A06"}),
    ("A06", {"vehicle_photo": "no_photo", "plate_photo": "yes"}, {"vehicle_photo": "yes"}, {"A05"}),
    ("A07", {"speed": "dispute", "speed_docs": "yes"}, {"speed": "no"}, set()),
    ("A08", {"place_time_match": "wrong_place", "place_time_docs": "yes"}, {"place_time_match": "yes"}, set()),
    ("A09", {"sign": "absent", "sign_docs": "yes"}, {"sign": "none"}, set()),
    ("A10", {"marking": "conflict", "marking_docs": "yes"}, {"marking": "none"}, set()),
)

BASE_NEGATIVE_LEGAL_ANSWERS = {
    "appeal_received_at": "28.08.2026",
    "complaint_recipient": "Тверской районный суд города Москвы",
    "correspondence_address": "Москва, Тверская улица, дом 1",
    "driver": "owner",
    "vehicle_photo": "yes",
    "plate_photo": "yes",
    "place_time_match": "yes",
    "speed": "no",
    "camera": "none",
    "sign": "none",
    "marking": "none",
    "owner_data_match": "yes",
    "previous_resolution": "no",
    "article_qualification": "no",
    "duplicate": "no",
    "emergency": "no",
}

@pytest.mark.parametrize(
    ("expected_code", "positive_answers", "negative_answers", "conflicting_codes"),
    FIRST_TEN_SCENARIO_CASES,
    ids=[case[0] for case in FIRST_TEN_SCENARIO_CASES],
)
def test_first_ten_scenarios_have_positive_and_negative_control_examples(
    expected_code: str,
    positive_answers: dict[str, str],
    negative_answers: dict[str, str],
    conflicting_codes: set[str],
) -> None:
    positive_results = evaluate_rules(
        {**BASE_NEGATIVE_LEGAL_ANSWERS, **positive_answers},
        notice_article="ч.2 ст.12.9 КоАП РФ",
    )
    negative_results = evaluate_rules(
        {**BASE_NEGATIVE_LEGAL_ANSWERS, **negative_answers},
        notice_article="ч.2 ст.12.9 КоАП РФ",
    )

    positive_codes = {result["code"] for result in positive_results}
    negative_codes = {result["code"] for result in negative_results}

    assert expected_code in positive_codes
    assert expected_code not in negative_codes
    assert positive_codes.isdisjoint(conflicting_codes)
    assert all(result["required_evidence"] for result in positive_results)
    assert all(result["evidence_items"] for result in positive_results)
    assert all(
        item["status"] in {status.value for status in EvidenceStatus}
        for result in positive_results
        for item in result["evidence_items"]
    )


def test_legal_questionnaire_branches_on_factual_answer() -> None:
    assert get_next_question({}).id == "appeal_received_at"
    answers = {
        "appeal_received_at": "28.08.2026",
        "complaint_recipient": "Тверской районный суд города Москвы",
        "correspondence_address": "Москва, Тверская улица, дом 1",
    }
    assert get_next_question({"appeal_received_at": "28.08.2026"}).id == (
        "complaint_recipient"
    )
    assert get_next_question(answers).id == "driver"
    assert get_next_question({**answers, "driver": "owner"}).id == "vehicle_photo"
    assert get_next_question({**answers, "driver": "other"}).id == "driver_docs"


def test_overdue_appeal_questions_are_asked_first() -> None:
    question = get_next_question({"appeal_received_at": "10.08.2026"})

    assert question is not None
    assert question.id == "appeal_delay_reason"
    assert get_next_question(
        {
            "appeal_received_at": "10.08.2026",
            "appeal_delay_reason": "late_receipt",
        }
    ).id == "complaint_recipient"


def test_non_speed_notice_skips_speed_questions_and_rule() -> None:
    answers = {
        "appeal_received_at": "28.08.2026",
        "complaint_recipient": "Тверской районный суд города Москвы",
        "correspondence_address": "Москва, Тверская улица, дом 1",
        "driver": "owner",
        "vehicle_photo": "yes",
        "plate_photo": "yes",
        "place_time_match": "yes",
    }

    next_question = get_next_question(answers, "ч.1 ст.12.16")
    results = evaluate_rules(
        {**answers, "speed": "dispute", "speed_docs": "yes"},
        "ч.1 ст.12.16",
    )

    assert next_question is not None
    assert next_question.id == "camera"
    assert all(result["code"] != "A07" for result in results)


def test_rules_version_changes_from_september_first() -> None:
    assert select_rules_version_for_date(None) == "2026-08-28"
    assert select_rules_version_for_date(__import__("datetime").date(2026, 8, 31)) == "2026-08-28"
    assert select_rules_version_for_date(__import__("datetime").date(2026, 9, 1)) == "2026-09-01"

def test_parse_date_accepts_ocr_datetime_strings_for_rules_version() -> None:
    from app.services.legal_rules import parse_date

    assert parse_date("31.08.2026 23:59") == date(2026, 8, 31)
    assert parse_date("01.09.2026 00:01") == date(2026, 9, 1)
    assert parse_date("2026-09-01T12:30:00") == date(2026, 9, 1)


@pytest.mark.asyncio
async def test_legal_assessment_persists_completed_result(
    db_session: AsyncSession,
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_020, None, "Legal", "Owner"
    )
    case = await CaseService(db_session).create(user.id)
    service = LegalAssessmentService(db_session)
    assessment = await service.start(case.id)
    answers = {
        "appeal_received_at": "28.08.2026",
        "complaint_recipient": "Тверской районный суд города Москвы",
        "correspondence_address": "Москва, Тверская улица, дом 1",
        "driver": "other",
        "driver_docs": "yes",
        "vehicle_photo": "yes",
        "plate_photo": "yes",
        "speed": "no",
        "camera": "none",
        "sign": "none",
        "place_time_match": "yes",
        "marking": "none",
        "owner_data_match": "yes",
        "previous_resolution": "no",
        "article_qualification": "no",
        "duplicate": "no",
        "emergency": "no",
    }

    while (question := get_next_question(assessment.answers)) is not None:
        await service.answer(assessment, question.id, answers[question.id])

    loaded = await service.get_for_case(case.id)
    assert loaded is not None
    assert loaded.status.value == "COMPLETED"
    assert [result["code"] for result in loaded.results] == ["A01"]
    assert loaded.completed_at is not None

    with pytest.raises(ValueError, match="already completed"):
        await service.answer(loaded, "driver", "owner")


@pytest.mark.asyncio
async def test_legal_assessment_saves_document_recipient_and_correspondence_address(
    db_session: AsyncSession,
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_021, None, "Legal", "Address"
    )
    case = await CaseService(db_session).create(user.id)
    service = LegalAssessmentService(db_session)
    assessment = await service.start(case.id)

    await service.answer(assessment, "appeal_received_at", "28.08.2026")
    next_question = await service.answer(
        assessment,
        "complaint_recipient",
        "  Тверской   районный суд города Москвы  ",
    )

    assert next_question is not None
    assert next_question.id == "correspondence_address"

    next_question = await service.answer(
        assessment,
        "correspondence_address",
        "  Москва,   Тверская улица, дом 1  ",
    )

    assert next_question is not None
    assert next_question.id == "driver"
    assert assessment.answers["complaint_recipient"] == (
        "Тверской районный суд города Москвы"
    )
    assert assessment.answers["correspondence_address"] == (
        "Москва, Тверская улица, дом 1"
    )


async def _create_completed_stage_three_case(
    db_session: AsyncSession,
) -> Case:
    user = await UserService(db_session).get_or_create(
        100_000_040, None, "Stage", "Four"
    )
    case = await CaseService(db_session).create(user.id)
    await RecognitionService(db_session).update_notice(
        case.id,
        FineNoticeFields(
            notice_number="18810177260901000123",
            notice_date="01.09.2026",
            fine_amount=1500,
            article="ч.2 ст.12.9",
            vehicle_plate="А000АА00",
            violation_datetime="01.09.2026 12:30",
            violation_place="Тестовая улица, дом 1",
            issuing_authority="Тестовый орган",
        ),
    )
    answers = {
        "appeal_received_at": "02.09.2026",
        "complaint_recipient": "Тверской районный суд города Москвы",
        "correspondence_address": "Москва, Тверская улица, дом 1",
        "driver": "owner",
        "vehicle_photo": "yes",
        "plate_photo": "yes",
        "place_time_match": "yes",
        "speed": "dispute",
        "speed_docs": "no",
        "camera": "calibration",
        "sign": "none",
        "marking": "none",
        "owner_data_match": "yes",
        "previous_resolution": "no",
        "article_qualification": "no",
        "duplicate": "no",
        "emergency": "no",
    }
    assessment = LegalAssessment(
        case_id=case.id,
        status=LegalAssessmentStatus.COMPLETED,
        answers=answers,
        results=evaluate_rules(answers, notice_article="ч.2 ст.12.9"),
        rules_version="2026-09-01",
        completed_at=datetime.now(UTC),
    )
    db_session.add(assessment)
    await db_session.flush()
    return case


@pytest.mark.asyncio
async def test_legal_analysis_filters_unknown_rules_sources_and_facts(
    db_session: AsyncSession,
) -> None:
    case = await _create_completed_stage_three_case(db_session)
    analysis = await LegalAnalysisService(
        db_session, FakeDeepSeekClient()
    ).analyze_case(case.id)

    assert analysis.status is LegalAnalysisStatus.PENDING_CONFIRMATION
    assert analysis.model == get_settings().deepseek_model
    assert [ground["id"] for ground in analysis.grounds] == ["ground-valid"]
    assert analysis.grounds[0]["legal_rule_ids"] == ["A12"]
    assert analysis.grounds[0]["source_ids"] == ["koap-rf"]
    assert "сведения о поверке комплекса" in analysis.missing_evidence
    assert analysis.input_summary["case"]["facts"]["vehicle_plate"] == "А000АА00"
    assert analysis.input_summary["facts"]["complaint_recipient"] == (
        "Тверской районный суд города Москвы"
    )


@pytest.mark.asyncio
async def test_legal_ground_confirmation_and_rejection_are_saved(
    db_session: AsyncSession,
) -> None:
    case = await _create_completed_stage_three_case(db_session)
    service = LegalAnalysisService(db_session, FakeDeepSeekClient())
    analysis = await service.analyze_case(case.id)

    analysis = await service.set_ground_status(
        case.id, "ground-valid", LegalGroundStatus.CONFIRMED
    )
    assert analysis.status is LegalAnalysisStatus.CONFIRMED
    assert analysis.grounds[0]["status"] == LegalGroundStatus.CONFIRMED.value

    analysis = await service.set_ground_status(
        case.id, "ground-valid", LegalGroundStatus.REJECTED
    )
    assert analysis.status is LegalAnalysisStatus.PENDING_CONFIRMATION
    assert analysis.grounds[0]["status"] == LegalGroundStatus.REJECTED.value


@pytest.mark.asyncio
async def test_document_generation_uses_only_confirmed_grounds_and_writes_files(
    db_session: AsyncSession, tmp_path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(generated_document_service_module, "BACKEND_ROOT", tmp_path)
    monkeypatch.setattr(
        generated_document_service_module,
        "GENERATED_ROOT",
        tmp_path / "storage" / "generated",
    )
    case = await _create_completed_stage_three_case(db_session)
    fake_client = FakeDeepSeekClient()
    service = LegalAnalysisService(db_session, fake_client)
    analysis = await service.analyze_case(case.id)
    analysis.grounds = [
        {**analysis.grounds[0], "status": LegalGroundStatus.CONFIRMED.value},
        {
            "id": "rejected",
            "title": "Отклоненное основание",
            "description": "Не должно попасть в документы.",
            "supporting_fact_ids": ["article"],
            "legal_rule_ids": ["A12"],
            "source_ids": ["koap-rf"],
            "missing_evidence": [],
            "assumptions": [],
            "recommended": False,
            "status": LegalGroundStatus.REJECTED.value,
        },
    ]
    await db_session.flush()

    documents = await service.generate_documents(case.id)

    assert len(documents) == 4
    assert {document.document_type for document in documents} == {
        GeneratedDocumentType.COMPLAINT,
        GeneratedDocumentType.EVIDENCE_PETITION,
    }
    assert all("Отклоненное основание" not in prompt for prompt in fake_client.document_prompts)

    docx_document = next(document for document in documents if document.original_filename.endswith(".docx"))
    pdf_document = next(document for document in documents if document.original_filename.endswith(".pdf"))
    docx_text = "\n".join(
        paragraph.text for paragraph in DocxDocument(tmp_path / docx_document.file_path).paragraphs
    )
    pdf = pymupdf.open(tmp_path / pdf_document.file_path)
    try:
        pdf_text = "\n".join(page.get_text() for page in pdf)
    finally:
        pdf.close()

    assert "Прошу отменить постановление" in docx_text
    assert "Прошу отменить постановление" in pdf_text
    assert "В Тверской районный суд города Москвы" in docx_text
    assert "В Тверской районный суд города Москвы" in pdf_text
    assert all(
        '"complaint_recipient": "Тверской районный суд города Москвы"' in prompt
        for prompt in fake_client.document_prompts
    )


@pytest.mark.asyncio
async def test_check_fine_button_does_not_restart_questionnaire(
    db_session: AsyncSession,
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_030, None, "Legal", "Flow"
    )
    case = await CaseService(db_session).create(user.id)
    assessment = await LegalAssessmentService(db_session).start(case.id)

    message = SimpleNamespace(
        from_user=SimpleNamespace(id=user.telegram_id),
        answer=AsyncMock(),
    )
    state = AsyncMock()
    state.get_state.return_value = LegalQuestionnaire.waiting_for_answer.state
    state.get_data.return_value = {
        "case_id": case.id,
        "question_id": "appeal_received_at",
    }

    before = await db_session.scalar(select(func.count(Case.id)).where(Case.user_id == user.id))

    await create_case(message, state, db_session)

    after = await db_session.scalar(select(func.count(Case.id)).where(Case.user_id == user.id))

    assert before == after == 1
    assert message.answer.await_count == 1
    assert "Сначала ответьте на текущий вопрос анкеты" in message.answer.await_args.args[0]
    assert assessment.status.value == "IN_PROGRESS"


@pytest.mark.asyncio
async def test_date_answer_uses_telegram_user_to_continue_questionnaire(
    db_session: AsyncSession,
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_031, None, "Legal", "Date"
    )
    case = await CaseService(db_session).create(user.id)
    assessment = await LegalAssessmentService(db_session).start(case.id)

    message = SimpleNamespace(
        from_user=SimpleNamespace(id=user.telegram_id),
        text="29.08.2026",
        answer=AsyncMock(),
    )
    state = AsyncMock()
    state.get_data.return_value = {
        "case_id": case.id,
        "question_id": "appeal_received_at",
    }

    await answer_date_question(message, state, db_session)

    assert assessment.answers["appeal_received_at"] == "29.08.2026"
    assert message.answer.await_count == 1
    assert "Дело не найдено" not in message.answer.await_args.args[0]
    assert "Юридическая анкета" in message.answer.await_args.args[0]




def test_fine_notice_extractor_reads_basic_fields() -> None:
    text = (
        "Постановление № 18810177230801000123 от 01.08.2026. "
        "УИН 18810177230801000123456. Штраф 1500 руб. "
        "ст. 12.9 КоАП РФ. Автомобиль А123ВС77. "
        "Место нарушения: Москва, Тверская улица, дом 1."
    )

    fields = FineNoticeExtractor().extract(text)

    assert fields.notice_number == "18810177230801000123"
    assert fields.notice_date == "01.08.2026"
    assert fields.uin == "18810177230801000123456"
    assert fields.fine_amount == 1500
    assert fields.article == "ст. 12.9"
    assert fields.vehicle_plate == "А123ВС77"


def test_fine_notice_extractor_reads_gosuslugi_format() -> None:
    text = """
    27.08.2026, 12:35
    Постановление Nº18810518260723020721 от 23.07.2026
    Сумма начисления
    3 000 ₽
    Госномер
    P225MO159
    Время и место нарушения
    18.07.2026 (Суббота), 07:31
    АВТОДОРОГА ВОТКИНСК - КЕЛЬЧИНО - ГР. ПЕРМСКОГО КРАЯ, KM 16+700
    Штраф выписан
    ЦАФАП в ОДД Госавтоинспекции МВД по Удмуртской Республике
    https://www.gosuslugi.ru/pay/uin/18810518260723020721
    """

    fields = FineNoticeExtractor().extract(text)

    assert fields.notice_number == "18810518260723020721"
    assert fields.notice_date == "23.07.2026"
    assert fields.uin == "18810518260723020721"
    assert fields.fine_amount == 3000
    assert fields.vehicle_plate == "Р225МО159"
    assert fields.violation_datetime == "18.07.2026 07:31"
    assert fields.violation_place is not None
    assert "АВТОДОРОГА ВОТКИНСК" in fields.violation_place
    assert fields.issuing_authority is not None
    assert "ЦАФАП" in fields.issuing_authority


def test_fine_notice_extractor_reads_copy_notice_format() -> None:
    text = """
    КОПИЯ ПОСТАНОВЛЕНИЕ 18810577260848002037
    по делу об административном правонарушении
    18810577260848002037
    17.08.2026
    ЦАФАП ОДД ГИБДД ГУ МВД России по г. Москве
    Я, инспектор, рассмотрев материалы.
    УСТАНОВИЛ:
    03.07.2026 в 17:52:15 по адресу ул.Мытная, д.18, г. Москва водитель,
    управляя транспортным средством, государственный регистрационный знак 0315УС797,
    нарушил требование.
    предусмотренного ч.1 ст.12.16 КоАП РФ, и назначить ему административное
    наказание в виде административного штрафа в размере 750 руб.
    УИН: 18810577260848002037.
    """

    fields = FineNoticeExtractor().extract(text)

    assert fields.notice_number == "18810577260848002037"
    assert fields.notice_date == "17.08.2026"
    assert fields.uin == "18810577260848002037"
    assert fields.fine_amount == 750
    assert fields.article == "ч.1 ст.12.16"
    assert fields.vehicle_plate == "0315УС797"
    assert fields.violation_datetime == "03.07.2026 17:52"
    assert fields.violation_place == "ул.Мытная, д.18, г. Москва"
    assert fields.issuing_authority == "ЦАФАП ОДД ГИБДД ГУ МВД России по г. Москве"


@pytest.mark.asyncio
async def test_recognition_service_creates_notice_and_updates_fields(
    db_session: AsyncSession,
) -> None:
    user = await UserService(db_session).get_or_create(
        100_000_004, None, "Recognition", "Owner"
    )
    case = await CaseService(db_session).create(user.id)
    document = await DocumentService(db_session).create(
        case=case,
        telegram_file_id="recognition-file-id",
        original_filename="notice.pdf",
        mime_type="application/pdf",
        local_path=None,
    )
    service = RecognitionService(db_session)

    recognition = await service.create_pending_for_document(case, document)
    notice = await service.update_notice(
        case.id,
        FineNoticeFields(
            notice_number="18810177230801000123",
            fine_amount=1500,
        ),
    )

    assert recognition.status is RecognitionStatus.VERIFIED
    assert case.status is CaseStatus.READY
    assert notice.case_id == case.id
    assert notice.recognition_id == recognition.id
    assert notice.notice_number == "18810177230801000123"
    assert notice.fine_amount == 1500

    with pytest.raises(ValueError, match="cannot be overwritten"):
        await service.process_document(
            case.id,
            document,
            b"new content",
            DisabledOcrProvider(),
        )

    assert notice.notice_number == "18810177230801000123"
    assert notice.fine_amount == 1500


@pytest.mark.asyncio
async def test_recognition_service_processes_document_with_ocr_provider(
    db_session: AsyncSession,
) -> None:
    class FakeOcrProvider:
        async def recognize(self, content: bytes, filename: str, mime_type: str | None):
            assert content == b"file-bytes"
            assert filename == "notice.jpg"
            assert mime_type == "image/jpeg"
            return OcrResult(
                text=(
                    "Постановление № 18810177230801000123. "
                    "УИН 18810177230801000123456. Штраф 1500 руб."
                )
            )

    user = await UserService(db_session).get_or_create(
        100_000_005, None, "Ocr", "Owner"
    )
    case = await CaseService(db_session).create(user.id)
    document = await DocumentService(db_session).create(
        case=case,
        telegram_file_id="ocr-file-id",
        original_filename="notice.jpg",
        mime_type="image/jpeg",
        local_path=None,
    )

    recognition = await RecognitionService(db_session).process_document(
        case.id,
        document,
        b"file-bytes",
        FakeOcrProvider(),
    )

    assert recognition.status is RecognitionStatus.RECOGNIZED
    assert case.status is CaseStatus.IN_PROGRESS
    assert recognition.raw_text is not None
    assert "Постановление" in recognition.raw_text
    loaded_case = await CaseService(db_session).get_for_user(user.id, case.id)
    assert loaded_case is not None
    case_card = _case_detail_text(loaded_case)
    assert "Распознавание: Распознано" in case_card
    assert "Номер постановления: 18810177230801000123" in case_card
    assert "Сумма штрафа: 1 500 руб." in case_card


@pytest.mark.asyncio
async def test_ocr_space_provider_rejects_oversized_file() -> None:
    settings = get_settings().model_copy(update={"ocr_max_file_size_bytes": 3})

    with pytest.raises(RuntimeError, match="превышает лимит"):
        await OcrSpaceProvider(settings).recognize(
            b"four", "notice.jpg", "image/jpeg"
        )


@pytest.mark.asyncio
async def test_ocr_space_provider_reads_successful_response(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    statuses = iter((503, 200))
    requests = 0

    class FakeResponse:
        def __init__(self, status: int):
            self.status = status

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        async def json(self, content_type=None):
            return {
                "IsErroredOnProcessing": False,
                "ParsedResults": [
                    {"ParsedText": "Постановление № 18810177230801000123"}
                ],
            }

    class FakeClientSession:
        def __init__(self, timeout):
            self.timeout = timeout

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        def post(self, endpoint, data, headers):
            nonlocal requests
            requests += 1
            assert endpoint == OcrSpaceProvider.endpoint
            assert headers["apikey"]
            return FakeResponse(next(statuses))

    async def no_sleep(delay: float) -> None:
        return None

    monkeypatch.setattr(
        ocr_service_module.aiohttp, "ClientSession", FakeClientSession
    )
    monkeypatch.setattr(ocr_service_module.asyncio, "sleep", no_sleep)
    settings = get_settings().model_copy(
        update={"ocr_max_file_size_bytes": 1_000_000}
    )

    result = await OcrSpaceProvider(settings).recognize(
        b"image", "notice.jpg", "image/jpeg"
    )

    assert result.text == "Постановление № 18810177230801000123"
    assert requests == 2


@pytest.mark.asyncio
async def test_ocr_space_provider_renders_oversized_pdf_pages(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    requests = 0

    class FakeResponse:
        status = 200

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        async def json(self, content_type=None):
            return {
                "IsErroredOnProcessing": False,
                "ParsedResults": [{"ParsedText": "Распознанная страница"}],
            }

    class FakeClientSession:
        def __init__(self, timeout):
            self.timeout = timeout

        async def __aenter__(self):
            return self

        async def __aexit__(self, *args):
            return None

        def post(self, endpoint, data, headers):
            nonlocal requests
            requests += 1
            return FakeResponse()

    document = pymupdf.open()
    for page_number in range(2):
        page = document.new_page()
        page.insert_text((72, 72), f"Synthetic notice page {page_number + 1}")
    content = document.tobytes() + b" " * 1_100_000
    document.close()

    monkeypatch.setattr(
        ocr_service_module.aiohttp, "ClientSession", FakeClientSession
    )
    settings = get_settings().model_copy(
        update={
            "ocr_max_file_size_bytes": 1_000_000,
            "ocr_max_pdf_pages": 3,
        }
    )

    result = await OcrSpaceProvider(settings).recognize(
        content, "notice.pdf", "application/pdf"
    )

    assert requests == 2
    assert result.text == "Распознанная страница\n\nРаспознанная страница"
