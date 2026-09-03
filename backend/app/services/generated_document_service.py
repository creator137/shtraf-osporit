from pathlib import Path
from textwrap import wrap
from uuid import uuid4

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_RIGHT
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models import (
    Case,
    GeneratedDocument,
    GeneratedDocumentFormat,
    GeneratedDocumentType,
    LegalAnalysis,
)
from app.services.document_service import BACKEND_ROOT


GENERATED_ROOT = BACKEND_ROOT / "storage" / "generated"


class GeneratedDocumentService:
    def __init__(self, session: AsyncSession) -> None:
        self.session = session

    async def save_pair(
        self,
        *,
        case: Case,
        analysis: LegalAnalysis,
        document_type: GeneratedDocumentType,
        addressee: str,
        title: str,
        sections: list[str],
    ) -> list[GeneratedDocument]:
        base_name = f"{document_type.value.lower()}-{uuid4().hex}"
        case_dir = GENERATED_ROOT / str(case.id)
        case_dir.mkdir(parents=True, exist_ok=True)
        docx_path = case_dir / f"{base_name}.docx"
        pdf_path = case_dir / f"{base_name}.pdf"

        _write_docx(docx_path, addressee, title, sections)
        _write_pdf(pdf_path, addressee, title, sections)

        created = [
            GeneratedDocument(
                case_id=case.id,
                legal_analysis_id=analysis.id,
                document_type=document_type,
                file_format=GeneratedDocumentFormat.DOCX,
                file_path=docx_path.relative_to(BACKEND_ROOT).as_posix(),
                original_filename=f"{document_type.value.lower()}.docx",
            ),
            GeneratedDocument(
                case_id=case.id,
                legal_analysis_id=analysis.id,
                document_type=document_type,
                file_format=GeneratedDocumentFormat.PDF,
                file_path=pdf_path.relative_to(BACKEND_ROOT).as_posix(),
                original_filename=f"{document_type.value.lower()}.pdf",
            ),
        ]
        self.session.add_all(created)
        await self.session.flush()
        return created


def generated_file_path(document: GeneratedDocument) -> Path:
    file_path = (BACKEND_ROOT / document.file_path).resolve()
    storage_root = (BACKEND_ROOT / "storage").resolve()
    if storage_root not in file_path.parents:
        raise ValueError("Generated document path is outside storage")
    return file_path


def _write_docx(
    path: Path, addressee: str, title: str, sections: list[str]
) -> None:
    document = DocxDocument()
    section = document.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(56)
    section.right_margin = Pt(56)
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    addressee_paragraph = document.add_paragraph(_addressee_line(addressee))
    addressee_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_heading(title, level=1)
    for section_text in sections:
        for paragraph in section_text.split("\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
    document.save(path)


def _write_pdf(
    path: Path, addressee: str, title: str, sections: list[str]
) -> None:
    font_name = _register_cyrillic_font()
    styles = getSampleStyleSheet()
    styles["Title"].fontName = font_name
    styles["Normal"].fontName = font_name
    styles["Normal"].fontSize = 11
    styles["Normal"].leading = 15
    addressee_style = styles["Normal"].clone("Addressee")
    addressee_style.alignment = TA_RIGHT

    story = [
        Paragraph(_escape_pdf_text(_addressee_line(addressee)), addressee_style),
        Spacer(1, 18),
        Paragraph(_escape_pdf_text(title), styles["Title"]),
        Spacer(1, 12),
    ]
    for section_text in sections:
        for paragraph in section_text.split("\n"):
            if paragraph.strip():
                wrapped = "<br/>".join(wrap(_escape_pdf_text(paragraph.strip()), 110))
                story.append(Paragraph(wrapped, styles["Normal"]))
                story.append(Spacer(1, 8))
    SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=56,
        leftMargin=56,
        topMargin=56,
        bottomMargin=56,
    ).build(story)


def _register_cyrillic_font() -> str:
    candidates = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    )
    for candidate in candidates:
        if Path(candidate).is_file():
            pdfmetrics.registerFont(TTFont("Stage4Cyrillic", candidate))
            return "Stage4Cyrillic"
    return "Helvetica"


def _escape_pdf_text(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _addressee_line(addressee: str) -> str:
    value = " ".join(addressee.split())
    return value if value.casefold().startswith("в ") else f"В {value}"
